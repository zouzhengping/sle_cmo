#!/usr/bin/env python3
"""
自动合并论文文档
按照PLOS ONE标准顺序合并所有章节
"""

from docx import Document
from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "docs/word_format"
OUT_FILE = DOC_DIR / "paper_merged_final.docx"

# 定义合并顺序
MERGE_ORDER = [
    "paper_title_page_and_statements.docx",  # Title, Authors, Statements
    "paper_abstract.docx",                    # Abstract
    "paper_introduction.docx",                # Introduction
    "paper_methods_results.docx",            # Methods and Results
    "paper_discussion.docx",                  # Discussion
    "paper_references.docx",                  # References
]

def merge_documents():
    """合并所有文档"""
    print("=" * 60)
    print("开始合并论文文档")
    print("=" * 60)
    
    # 创建新文档
    merged_doc = Document()
    
    for i, filename in enumerate(MERGE_ORDER, 1):
        filepath = DOC_DIR / filename
        
        if not filepath.exists():
            print(f"⚠ 警告: 文件不存在 - {filename}")
            continue
        
        print(f"\n[{i}/{len(MERGE_ORDER)}] 正在合并: {filename}")
        
        try:
            # 读取文档
            source_doc = Document(filepath)
            
            # 复制所有段落
            for para in source_doc.paragraphs:
                new_para = merged_doc.add_paragraph()
                new_para.style = para.style
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
            
            # 复制所有表格
            for table in source_doc.tables:
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = table.style
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
            
            # 添加分页符（除了最后一个文件）
            if i < len(MERGE_ORDER):
                merged_doc.add_page_break()
            
            print(f"  ✓ 成功合并")
            
        except Exception as e:
            print(f"  ✗ 错误: {e}")
            continue
    
    # 保存合并后的文档
    merged_doc.save(OUT_FILE)
    print(f"\n{'='*60}")
    print(f"✅ 合并完成！")
    print(f"输出文件: {OUT_FILE}")
    print(f"{'='*60}")
    
    # 提示
    print("\n📌 下一步操作:")
    print("1. 打开合并后的文档检查格式")
    print("2. 插入所有图片到相应位置")
    print("3. 添加Figure Legends")
    print("4. 检查页码和格式")
    print("\n详细指南: docs/manuscript_merge_guide.md")

if __name__ == "__main__":
    merge_documents()



